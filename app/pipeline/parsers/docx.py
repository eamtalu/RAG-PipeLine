"""Stage 3b — DOCX parser using python-docx."""

import io

from docx import Document

from app.models.document import ParsedDocument, HeadingNode
from app.pipeline.parsers.base import BaseParser

_HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4}


class DocxParser(BaseParser):
    def parse(self, data: bytes, filename: str, mime_type: str) -> ParsedDocument:
        doc = Document(io.BytesIO(data))

        paragraphs: list[str] = []
        headings: list[HeadingNode] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            paragraphs.append(text)

            level = _HEADING_STYLES.get(para.style.name)  # type: ignore[union-attr]
            if level is not None:
                headings.append(HeadingNode(level=level, title=text))

        raw_text = "\n\n".join(paragraphs)
        title = headings[0].title if headings else None

        return ParsedDocument(
            source_filename=filename,
            mime_type=mime_type,
            title=title,
            raw_text=raw_text,
            headings=headings,
        )
